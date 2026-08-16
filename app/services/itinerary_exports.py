import html
import io
from dataclasses import dataclass
from datetime import date

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy.orm import Session

from app.models.itinerary import Itinerary
from app.services.meituan_itinerary_parser import category_label


MAIN_HEADERS = ("日期", "城市", "时间", "游玩地点/项目", "交通", "重点内容", "费用估算（元）")


@dataclass(frozen=True)
class ItineraryExportFile:
    content: bytes
    media_type: str
    filename: str


class ItineraryExportService:
    """Export only confirmed decisions as a directly actionable itinerary."""

    def __init__(self, db: Session) -> None:
        self._db = db

    def export_html(self, itinerary: Itinerary) -> ItineraryExportFile:
        rows = self._main_rows(itinerary)
        rendered = "".join("<tr>" + "".join(f"<td>{html.escape(str(value if value is not None else ''))}</td>" for value in row) + "</tr>" for row in rows)
        content = f"""<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'><title>{html.escape(itinerary.title)}</title><style>body{{font-family:'Microsoft YaHei',sans-serif;max-width:1200px;margin:auto;padding:30px}}table{{border-collapse:collapse;width:100%}}th,td{{border:1px solid #c8d5ce;padding:9px;vertical-align:top;text-align:left}}th{{background:#eaf2ed}}.price-notice{{color:#a35f14;font-weight:700}}</style></head><body><h1>{html.escape(itinerary.title)}</h1><p class='price-notice'>价格仅作参考，请以实际情况为准。</p><table><thead><tr>{''.join(f'<th>{header}</th>' for header in MAIN_HEADERS)}</tr></thead><tbody>{rendered}</tbody></table></body></html>"""
        return ItineraryExportFile(content.encode("utf-8"), "text/html; charset=utf-8", self._filename(itinerary, "html"))

    def export_docx(self, itinerary: Itinerary) -> ItineraryExportFile:
        document = Document()
        document.styles["Normal"].font.name = "Microsoft YaHei"
        document.styles["Normal"].font.size = Pt(10)
        document.add_heading(itinerary.title, 0)
        table = document.add_table(rows=1, cols=len(MAIN_HEADERS)); table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, MAIN_HEADERS): cell.text = header
        for row in self._main_rows(itinerary):
            for cell, value in zip(table.add_row().cells, row): cell.text = str(value if value is not None else "")
        output = io.BytesIO(); document.save(output)
        return ItineraryExportFile(output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", self._filename(itinerary, "docx"))

    def export_xlsx(self, itinerary: Itinerary) -> ItineraryExportFile:
        workbook = Workbook(); sheet = workbook.active; sheet.title = "行程计划"
        sheet.append(("价格仅作参考，请以实际情况为准。",))
        sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(MAIN_HEADERS))
        sheet.append(MAIN_HEADERS)
        for row in self._main_rows(itinerary): sheet.append(row)
        if sheet.max_row >= 4:
            sheet.cell(sheet.max_row, 7).value = f"=SUM(G3:G{sheet.max_row - 1})"
        self._format_main_sheet(sheet)
        self._append_reference_sheet(workbook, itinerary)
        output = io.BytesIO(); workbook.save(output)
        return ItineraryExportFile(output.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", self._filename(itinerary, "xlsx"))

    def _main_rows(self, itinerary: Itinerary) -> list[tuple[object, ...]]:
        rows: list[tuple[object, ...]] = []
        for day in sorted(itinerary.days, key=lambda value: value.day_number):
            context, city, date_text, first = day.travel_context or {}, (day.travel_context or {}).get("city") or itinerary.destination_city or "待确认", day.itinerary_date.isoformat(), True
            # A Skill-normalized itinerary is already an ordered daily plan.
            # Export it first and omit the duplicate IP anchor rows.
            meituan_items = context.get("meituan_items") or []
            if meituan_items:
                for item in meituan_items:
                    time_slot = item.get("time_slot") or "待确认"
                    end_time = item.get("end_time")
                    note = self._display_note(item.get("note"))
                    if item.get("price_missing"):
                        note = f"{note}；费用待确认，未计入总预算。"
                    rows.append((
                        date_text if first else "",
                        item.get("city") or city,
                        f"{time_slot}–{end_time}" if end_time else time_slot,
                        item.get("name"),
                        item.get("transport") or "待确认",
                        f"{category_label(item.get('category'))} · {note}",
                        None if item.get("price_missing") else item.get("price", 0),
                    ))
                    first = False
                continue
            for leg in context.get("intercity_transport") or []:
                rows.append((date_text if first else "", city, "按已确认班次", f"{leg['label']}：{leg['from']} → {leg['to']}", self._leg_transport(leg), "用户已确认大交通", leg.get("price", self._leg_estimate(leg))))
                first = False
            for stop in sorted(day.stops, key=lambda value: value.stop_order):
                price = (context.get("landmark_costs") or {}).get(str(stop.landmark_id), 0)
                rows.append((date_text if first else "", city, f"{stop.time_slot}–{self._end_time(stop.time_slot, stop.planned_minutes)}", stop.landmark.name, self._local_transport(stop.landmark.transit_text, itinerary.transport_preference), stop.selection_reason, price))
                first = False
            for item in day.supplemental_items or []:
                time_slot = item.get("time_slot") or "按确认安排"
                end_time = item.get("end_time") or self._end_time(time_slot, item.get("planned_minutes", 0))
                transport = item.get("transport") or "按当天动线"
                note = item.get("note") or "用户已确认补充景点"
                if item.get("price_missing"):
                    note = f"{note}；费用待确认，未计入总预算。"
                rows.append((date_text if first else "", item.get("city") or city, f"{time_slot}–{end_time}" if item.get("time_slot") else time_slot, item.get("name"), transport, note, item.get("price", 0)))
                first = False
            for food in context.get("confirmed_food_events") or []:
                end_time = food.get("end_time") or self._end_time(food["time_slot"], food.get("planned_minutes", 0))
                rows.append((date_text if first else "", city, f"{food['time_slot']}–{end_time}", f"用餐：{food['name']}", "按当天动线", food.get("note") or "用户已确认餐饮", food.get("price", 0)))
                first = False
            lodging = context.get("lodging") or {}
            if lodging.get("mode") == "confirmed" and lodging.get("overnight", True):
                rows.append((date_text if first else "", city, "晚间", f"入住：{lodging['name']}", "返回已确认住宿", lodging.get("address") or "用户已确认住宿", lodging.get("reference_price_per_night", 0)))
        total = sum(float(row[6] or 0) for row in rows)
        rows.append(("", "", "", "总预算", "", "已按表内已确认项目汇总", total))
        return rows

    @staticmethod
    def _end_time(start: str, minutes: int) -> str:
        try:
            hour, minute = (int(value) for value in start.split(":", 1)); total = hour * 60 + minute + minutes
            return f"{total // 60:02d}:{total % 60:02d}"
        except (TypeError, ValueError): return "结束时间待确认"

    @staticmethod
    def _display_note(value: str | None) -> str:
        """Do not expose old supplier labels retained in historical records."""
        note = value or "行程安排"
        if note.startswith("美团酒旅推荐："):
            return "车次：" + note.removeprefix("美团酒旅推荐：")
        if note.startswith("美团酒旅推荐住宿；"):
            return note.removeprefix("美团酒旅推荐住宿；") or "住宿安排"
        return note

    @staticmethod
    def _local_transport(transit: str | None, preference: str | None) -> str:
        return transit or f"{preference or '打车'}；1 公里以内步行"

    @staticmethod
    def _leg_transport(leg: dict) -> str:
        return ("飞机" if leg.get("mode") == "flight" else "高铁/火车") + f"：{leg.get('option_id') or '待确认'}" + (f"（{leg['seat']}）" if leg.get("seat") else "")

    @staticmethod
    def _leg_estimate(leg: dict) -> float | None:
        prices = leg.get("train_prices") or leg.get("flight_prices") or []
        return min(prices) if prices else None

    def _append_reference_sheet(self, workbook: Workbook, itinerary: Itinerary) -> None:
        sheet = workbook.create_sheet("参考说明")
        budget = itinerary.budget_summary or {}
        sheet.append(("项目", "内容")); sheet.append(("预算上限（元）", itinerary.budget_amount or "未填写")); sheet.append(("确认项目合计（元）", budget.get("estimated_amount") or 0)); sheet.append(("核算口径", budget.get("status") or "仅汇总表内项目"))
        for key, label in (("transport", "交通"), ("lodging", "住宿"), ("scenic", "景点/门票"), ("food", "餐饮"), ("other", "其他")):
            if key in (budget.get("breakdown") or {}):
                sheet.append((f"{label}小计（元）", budget["breakdown"][key]))
        if budget.get("unpriced_item_count"):
            sheet.append(("未标价项目数", budget["unpriced_item_count"]))
        sheet.column_dimensions["A"].width = 22; sheet.column_dimensions["B"].width = 70
        for row in sheet.iter_rows():
            for cell in row: cell.alignment = Alignment(vertical="top", wrap_text=True)

    @staticmethod
    def _format_main_sheet(sheet) -> None:
        sheet["A1"].font = Font(bold=True, color="A35F14"); sheet["A1"].alignment = Alignment(horizontal="left", vertical="center")
        for cell in sheet[2]: cell.font = Font(bold=True); cell.fill = PatternFill("solid", fgColor="EAF2ED"); cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in sheet.iter_rows(min_row=3):
            for cell in row: cell.alignment = Alignment(vertical="top", wrap_text=True)
        for index, width in enumerate((16, 14, 18, 34, 36, 52, 16), start=1): sheet.column_dimensions[chr(64 + index)].width = width
        sheet.freeze_panes = "A3"; sheet.auto_filter.ref = f"A2:G{sheet.max_row}"

    @staticmethod
    def _filename(itinerary: Itinerary, extension: str) -> str:
        return f"itinerary-{itinerary.id}-{date.today().isoformat()}.{extension}"
